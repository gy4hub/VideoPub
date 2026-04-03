"""元数据解析（JSON / Word / PDF → PublishTask）"""

import json
import re
from datetime import datetime
from pathlib import Path

from videopub.core.models import Platform, PlatformMeta, PublishTask


# ── 平台名称映射（中文/英文 → Platform enum）────────────────────────────────
PLATFORM_ALIASES: dict[str, Platform] = {
    "wechat": Platform.WECHAT,
    "视频号": Platform.WECHAT,
    "微信视频号": Platform.WECHAT,
    "douyin": Platform.DOUYIN,
    "抖音": Platform.DOUYIN,
    "bilibili": Platform.BILIBILI,
    "b站": Platform.BILIBILI,
    "youtube": Platform.YOUTUBE,
}

# ── 字段名映射（中文/英文 → 内部字段名）──────────────────────────────────────
FIELD_ALIASES: dict[str, str] = {
    # 通用
    "视频路径": "video_path",
    "video_path": "video_path",
    "封面路径": "cover_path",
    "cover_path": "cover_path",
    # 平台元字段
    "标题": "title",
    "title": "title",
    "短标题": "short_title",
    "short_title": "short_title",
    "简介": "description",
    "描述": "description",
    "description": "description",
    "标签": "tags",
    "tags": "tags",
    "分类": "category",
    "category": "category",
    "首评": "first_comment",
    "first_comment": "first_comment",
    "是否原创": "is_original",
    "原创": "is_original",
    "is_original": "is_original",
    "定时发布": "scheduled_time",
    "scheduled_time": "scheduled_time",
}


# ── 文件查找 ─────────────────────────────────────────────────────────────────

def _find_metadata_file(folder_path: Path) -> Path:
    """在文件夹中查找元数据文件，优先级: .json > .docx > .pdf"""
    for ext in (".json", ".docx", ".pdf"):
        candidates = sorted(folder_path.glob(f"metadata{ext}"))
        if candidates:
            return candidates[0]

    # 也支持任意名字的 json 文件，但明确排除隐藏文件（如 .status 目录下的文件虽然 glob("*.json") 不递归找不到，
    # 但如果是顶级目录下的隐藏文件，需要排除）
    json_files = [f for f in sorted(folder_path.glob("*.json")) if not f.name.startswith(".")]
    if json_files:
        return json_files[0]

    raise FileNotFoundError(f"在 {folder_path} 中未找到元数据文件 (metadata.json/.docx/.pdf)")


def _find_video_file(folder_path: Path) -> Path:
    """在文件夹中查找视频文件"""
    for ext in (".mp4", ".mov", ".avi", ".mkv"):
        candidates = sorted(folder_path.glob(f"*{ext}"))
        if candidates:
            return candidates[0]

    raise FileNotFoundError(f"在 {folder_path} 中未找到视频文件")


def _find_cover_file(folder_path: Path) -> Path:
    """在文件夹中查找封面图"""
    for ext in (".jpg", ".jpeg", ".png", ".webp"):
        candidates = sorted(folder_path.glob(f"*{ext}"))
        if candidates:
            return candidates[0]

    raise FileNotFoundError(f"在 {folder_path} 中未找到封面图")


# ── 工具函数 ──────────────────────────────────────────────────────────────────

def _parse_platform(platform_value: str) -> Platform:
    """将平台字符串解析为 Platform enum"""
    normalized = platform_value.strip().lower()
    alias_platform = PLATFORM_ALIASES.get(normalized)
    if alias_platform is not None:
        return alias_platform
    # 尝试中文 key 直接匹配
    for key, val in PLATFORM_ALIASES.items():
        if key == platform_value.strip():
            return val
    return Platform(normalized)


def _resolve_asset_path(
    folder_path: Path,
    raw_path: str | None,
    fallback_finder,
) -> Path:
    """解析资源文件路径，不存在时回退到自动查找"""
    if raw_path:
        candidate = Path(raw_path).expanduser()
        if not candidate.is_absolute():
            candidate = folder_path / candidate
        if candidate.exists():
            return candidate

    return fallback_finder(folder_path)


def _normalize_field(key: str) -> str | None:
    """将中文/英文 key 统一成内部字段名，找不到则返回 None"""
    stripped = key.strip()
    return FIELD_ALIASES.get(stripped) or FIELD_ALIASES.get(stripped.lower())


def _cast_field_value(field: str, raw: str):
    """将字段值从字符串转换为合适的类型"""
    raw = raw.strip()
    if field == "tags":
        return [t.strip() for t in re.split(r"[,，、]", raw) if t.strip()]
    if field == "is_original":
        return raw.lower() not in ("false", "否", "0", "no")
    if field == "scheduled_time":
        for fmt in (
            "%Y-%m-%d %H:%M:%S",
            "%Y-%m-%d %H:%M",
            "%Y-%m-%dT%H:%M:%S",
            "%Y-%m-%dT%H:%M",
            "%Y-%m-%dT%H:%M:%SZ",
        ):
            try:
                return datetime.strptime(raw, fmt)
            except ValueError:
                continue
        return None
    return raw


# ── JSON 解析 ─────────────────────────────────────────────────────────────────

def _parse_json(file_path: Path, folder_path: Path) -> PublishTask:
    """解析 JSON 格式的元数据文件"""
    with file_path.open("r", encoding="utf-8") as file:
        data = json.load(file)

    video_path = _resolve_asset_path(folder_path, data.get("video_path"), _find_video_file)
    cover_path = _resolve_asset_path(folder_path, data.get("cover_path"), _find_cover_file)

    platforms: list[PlatformMeta] = []
    for platform_data in data.get("platforms", []):
        platform_meta = PlatformMeta(
            platform=_parse_platform(platform_data["platform"]),
            title=platform_data["title"],
            short_title=platform_data.get("short_title"),
            description=platform_data["description"],
            tags=platform_data.get("tags", []),
            first_comment=platform_data.get("first_comment"),
            is_original=platform_data.get("is_original", True),
            scheduled_time=platform_data.get("scheduled_time"),
            category=platform_data.get("category"),
            extra=platform_data.get("extra", {}),
        )
        platforms.append(platform_meta)

    return PublishTask(
        video_path=video_path,
        cover_path=cover_path,
        platforms=platforms,
    )


# ── 通用文本解析（供 docx / pdf 共用）────────────────────────────────────────

def _parse_kv_text(text: str, folder_path: Path) -> PublishTask:
    """解析 key: value 格式的纯文本，支持 [platform] 节头。

    格式示例::

        视频路径: video.mp4
        封面路径: cover.jpg

        [bilibili]
        标题: 磁珠纯化误区
        简介: 本期视频...
        标签: 磁珠, 纯化, 科研
        分类: 201

        [youtube]
        标题: Magnetic Bead Mistakes
        描述: In this video...
        标签: magnetic bead, purification
        定时发布: 2026-04-03 20:00
    """
    global_data: dict = {}
    platforms: list[PlatformMeta] = []
    current_platform_raw: str | None = None
    current_fields: dict = {}

    def _flush_platform():
        nonlocal current_platform_raw, current_fields
        if current_platform_raw and current_fields:
            try:
                plat = _parse_platform(current_platform_raw)
                # 必填字段默认值
                meta = PlatformMeta(
                    platform=plat,
                    title=current_fields.get("title", ""),
                    short_title=current_fields.get("short_title"),
                    description=current_fields.get("description", ""),
                    tags=current_fields.get("tags", []),
                    first_comment=current_fields.get("first_comment"),
                    is_original=current_fields.get("is_original", True),
                    scheduled_time=current_fields.get("scheduled_time"),
                    category=current_fields.get("category"),
                )
                platforms.append(meta)
            except Exception:
                pass
        current_platform_raw = None
        current_fields = {}

    # 逐行解析
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#"):
            continue

        # [platform] 节头
        m_header = re.match(r"^\[(.+?)\]$", stripped)
        if m_header:
            _flush_platform()
            current_platform_raw = m_header.group(1).strip()
            current_fields = {}
            continue

        # key: value
        if ":" in stripped:
            key, _, value = stripped.partition(":")
            field = _normalize_field(key)
            if field is None:
                continue
            casted = _cast_field_value(field, value)

            if current_platform_raw is None:
                # 全局字段（video_path / cover_path）
                global_data[field] = casted
            else:
                current_fields[field] = casted

    _flush_platform()

    video_path = _resolve_asset_path(folder_path, global_data.get("video_path"), _find_video_file)
    cover_path = _resolve_asset_path(folder_path, global_data.get("cover_path"), _find_cover_file)

    return PublishTask(
        video_path=video_path,
        cover_path=cover_path,
        platforms=platforms,
    )


# ── Word 解析 ─────────────────────────────────────────────────────────────────

def _parse_docx(file_path: Path, folder_path: Path) -> PublishTask:
    """解析 Word (.docx) 格式的元数据文件"""
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ImportError("需要 python-docx: pip install python-docx") from exc

    doc = docx.Document(str(file_path))
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # 表格内容也提取（兼容表格格式）
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                lines.append(f"{cells[0]}: {cells[1]}")

    return _parse_kv_text("\n".join(lines), folder_path)


# ── PDF 解析 ──────────────────────────────────────────────────────────────────

def _parse_pdf(file_path: Path, folder_path: Path) -> PublishTask:
    """解析 PDF 格式的元数据文件"""
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError("需要 pdfplumber: pip install pdfplumber") from exc

    lines = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)

    return _parse_kv_text("\n".join(lines), folder_path)


# ── 公开入口 ──────────────────────────────────────────────────────────────────

def parse(folder_path: Path) -> PublishTask:
    """解析文件夹中的元数据文件，返回 PublishTask

    支持格式: JSON (.json) / Word (.docx) / PDF (.pdf)
    """
    folder_path = Path(folder_path)
    if not folder_path.is_dir():
        raise ValueError(f"路径不是文件夹: {folder_path}")

    meta_file = _find_metadata_file(folder_path)

    if meta_file.suffix == ".json":
        return _parse_json(meta_file, folder_path)
    if meta_file.suffix == ".docx":
        return _parse_docx(meta_file, folder_path)
    if meta_file.suffix == ".pdf":
        return _parse_pdf(meta_file, folder_path)
    raise ValueError(f"不支持的元数据格式: {meta_file.suffix}")
